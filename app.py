import streamlit as st
import pandas as pd
import json
from groq import Groq
from dotenv import load_dotenv
from graphviz import Digraph
import os, docx
from streamlit_lottie import st_lottie
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

                                                                # ==============================
                                                                # -----------SETUP--------------
                                                                # ==============================
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")
# COLORS
KMG_NAVY = RGBColor(31, 73, 125)
BLACK = RGBColor(0, 0, 0)
TEXT_GREY = RGBColor(80, 80, 80)
# INPUT SIZE LIMIT
MAX_CHARS = 15000
MAX_FILE_MB = 5
MAX_FILE_BYTES = MAX_FILE_MB * 1024 * 1024
                                                                # ==============================
                                                                # ------HELPER FUNCTIONS--------
                                                                # ==============================


with open("Assets/Web_Logo.json") as f:
    lottie = json.load(f)                                                    


def load_config():
    CONFIG_PATH = "Input/Config.xlsx"
    branding = pd.read_excel(CONFIG_PATH, sheet_name="BRANDING")
    contacts = pd.read_excel(CONFIG_PATH, sheet_name="CONTACTS")
    company = pd.read_excel(CONFIG_PATH, sheet_name="COMPANY")
    clients = pd.read_excel(CONFIG_PATH, sheet_name="CLIENT")
    branding_dict = branding.set_index("Client Code").to_dict("index")
    contacts_dict = contacts.set_index("Role").to_dict("index")
    company_dict = dict(zip(company["Company Fields"], company["Company Values"]))
    clients_dict = clients.set_index("Client Code").to_dict("index")
    return branding_dict, contacts_dict, company_dict, clients_dict

def extract_text_from_docx(file):
    doc = docx.Document(file)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return full_text[:15000]

def set_font(run, name="Trebuchet MS", size=11, color=None, bold=False, italic=False, underline=False):

    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color: run.font.color.rgb = color

def get_smart_flow_data(client, process_details):
    process_details = process_details[:4000]
    prompt = (f"Analyze this process: {process_details}. "
              f"Break it into a structured flowchart with actions and decisions. "
              f"Return ONLY a JSON object with this structure: "
              f"{{\"nodes\": [ {{\"id\": \"1\", \"label\": \"Step Name\", \"type\": \"action/decision\"}} ], "
              f"\"edges\": [ {{\"from\": \"1\", \"to\": \"2\", \"label\": \"Yes/No (optional)\"}} ] }}")
    try:
        completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.1-8b-instant",
            response_format={ "type": "json_object" }
        )
        return json.loads(completion.choices[0].message.content)
    except Exception:
        return {"nodes": [{"id":"1", "label":"Start", "type":"action"}], "edges": []}
    
 
def generate_smart_flowchart(data, output_path="flowchart"):
    dot = Digraph(comment='Process Flow', engine='dot')
    dot.attr(dpi='300')
    dot.attr(rankdir='TB', size='7,7!', ratio='fill')
    dot.attr(nodesep='0.5', ranksep='0.4', splines='polyline')
    dot.attr('node', fontname='Arial', fontsize='10', shape='rect',
             style='filled, rounded', color='#000000', fillcolor='#E3F2FD',
             width='2.0', height='0.6', penwidth='1.2')
 
    for node in data.get('nodes', []):
        label = node.get('label', '')
        if len(label) > 15:
            label = label.replace(' ', '\n', 1)
 
        node_type = node.get('type', '').lower()
        if node_type in ['start', 'end']:
            dot.node(node['id'], label, shape='capsule', fillcolor='#4285F4',
                     fontcolor='white', color='#000000', style='filled')
        elif node_type == 'decision':
            dot.node(node['id'], label, shape='diamond', fillcolor='#4285F4',
                     fontcolor='white', color='#000000', style='filled', width='1.4', height='0.9')
        else:
            dot.node(node['id'], label)
    for edge in data.get('edges', []):
        label_text = edge.get('label', '').strip()
       
        if label_text:
            dot.edge(edge['from'], edge['to'], taillabel=f" {label_text} ",
                     labelangle='-45', labeldistance='2.5', 
                     fontname='Arial', fontsize='9', fontcolor='#2c3e50')
        else:
            dot.edge(edge['from'], edge['to'], penwidth='1.0')
 
    dot.render(output_path, format='png', cleanup=True)
    return f"{output_path}.png"

def get_short_context(client, process_details):
    process_details = process_details[:12000]
    if not client:
        return process_details[:3000]

    prompt = f"""
    Summarize the following process in a concise way.
    Keep only key steps, systems used, inputs, outputs and decision points.

    Process:
    {process_details}
    """

    completion = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.1-8b-instant"
    )

    return completion.choices[0].message.content

def generate_ai_content(client, section_name, process_details, context_title):

    if not client:
        return "[API Key Missing]"

    section_prompts = {

        "INTRODUCTION":
        f"""
        Provide a brief introduction for the process '{context_title}'.

        Write in paragraph form covering:
        - business context
        - what the process does
        - systems involved

        Do not create any sub-headings.
        """,

        "AUDIENCE":
        f"""
        Identify the intended audience for this document.

        Start with one short paragraph explaining who should use this document,
        followed by a bullet list of roles.

        Do not add sub-headings.
        """,

        "PURPOSE":
        f"""
        Explain the purpose of this document in paragraph form.

        Cover:
        - why this document is created
        - how it will be used

        Do not create document objective or any sub-headings.
        """,

        "SCOPE":
        f"""
        Define the scope of the process '{context_title}' in paragraph form.

        Then provide the following as a list, each item on a new line:
        - in scope
        - out of scope
        - start point
        - end point

        Do not create sub-headings.
        """
    }

    prompt = f"""
    Context:
    {process_details}

    {section_prompts.get(section_name, "")}

    Write in a professional and technical tone.

    Do not generate sub-headings.
    Do not repeat the section title inside the content.
    Write the entire content in well-structured paragraphs.

    Where a list is required, return each item on a new line.
    Do NOT add *, -, or numbers.
    Do NOT add manual bullets.
    For all other text, use paragraph format.

    Return plain text only.
    No markdown.
    No asterisks.
    No additional headings.
    """

    try:
        completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
        )

        return completion.choices[0].message.content

    except Exception as e:
        return f"[AI Error: {str(e)}]"

def insert_constant_header(document, title, client_name, date_str, logo_path, client_cfg):

    section = document.sections[-1]
    header = section.header

    # usable page width
    usable_width = section.page_width - section.left_margin - section.right_margin

    # create table in header (WIDTH MANDATORY)
    htab = header.add_table(rows=4, cols=3, width=usable_width)
    htab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    htab.autofit = False
    htab.style = "Table Grid"

    from docx.shared import Emu
    htab.columns[0].width = Emu(int(usable_width * 0.5))
    htab.columns[1].width = Emu(int(usable_width * 0.3))
    htab.columns[2].width = Emu(int(usable_width * 0.2))

    def fill_h(row, col, text, is_bold=False):

        cell = htab.cell(row, col)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        htab.rows[row].height = Pt(12)  
        r = p.add_run(text)
        set_font(r, name="Calibri", size=10, color=TEXT_GREY, bold=is_bold)
    title_cell = htab.cell(0, 0).merge(htab.cell(0, 1))
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Process Flow Document for {title}")
    set_font(run, name="Calibri", size=10, color=TEXT_GREY, bold=True)
    fill_h(1, 0, f"Date: {date_str}")
    fill_h(2, 0, "Version: 0.1")
    fill_h(3, 0, f"Document Owner: {client_cfg['Document Owner']}")
    fill_h(1, 1, f"Classification: {client_cfg['Classification']}")
    fill_h(2, 1, f"Circulation: {client_cfg['Circulation']}")
    fill_h(3, 1, f"Client: {client_cfg['Client Name']}")

                                                            # ==============================
                                                            # --- LOGO CENTERING LOGIC ---
                                                            # ==============================

    logo_cell = htab.cell(0, 2).merge(htab.cell(3, 2))
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER # Vertical center

    if os.path.exists(logo_path):
        lp = logo_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER # Horizontal center
        lp.add_run().add_picture(logo_path, width=Inches(1.1))
    header.add_paragraph().paragraph_format.space_after = Pt(12)

                                                            # ==============================
                                                            # ------------UI----------------
                                                            # ==============================

st.set_page_config(page_title="ProcessCraft AI – PDD Studio", layout="wide")
col1, col2 = st.columns([1.2, 3])

col1, col2 = st.columns([1.2, 3])

with col1:
    if lottie:
        st_lottie(lottie, height=140)

with col2:
    st.title("ProcessCraft AI – PDD Studio")
    st.caption("AI-Powered Process Design & Documentation Platform")

st.markdown("""
<style>

/* Make uploader horizontal */
[data-testid="stFileUploader"] section {
    display: flex;
    align-items: center;
    justify-content: flex-start;
}

/* ICON + TEXT */
[data-testid="stFileUploader"] section::before {
    content: "  Drag and drop file here";
    font-weight: 500;
    color: #6b6b6b;
    display: flex;
    align-items: center;
    gap: 8px;

    background-image: url("https://cdn-icons-png.flaticon.com/512/109/109612.png");
    background-size: 22px 22px;
    background-repeat: no-repeat;
    background-position: left center;
    padding-left: 30px;
}

/* ❌ Hide Browse button */
[data-testid="stFileUploader"] button {
    display: none !important;
}

/* ❌ Hide default text */
[data-testid="stFileUploader"] section div span {
    display: none !important;
}

/* ❌ Hide helper text */
[data-testid="stFileUploader"] section div:nth-child(2) {
    display: none !important;
}

</style>
""", unsafe_allow_html=True)
client = Groq(api_key=api_key) if api_key else None
uploaded_file = st.file_uploader(
    f"Upload Source Process Input File (Maximum {MAX_FILE_MB} MB)",
    type=["txt", "docx"]
)
manual_input = st.text_area(
    f"Or Enter Process Details Manually (Maximum {MAX_CHARS:,} characters)",
    height=150
)
process_context = st.session_state.get("process_context", "")
manual_exceeded = False
file_exceeded = False
char_count = len(manual_input)
st.caption(f"Characters used: {char_count:,} / {MAX_CHARS:,}")
# -------- MANUAL INPUT CHECK --------
if char_count > MAX_CHARS:
    manual_exceeded = True
    st.error("Input limit exceeded. Please reduce the content to proceed.")



# -------- FILE CHECK --------
if uploaded_file:

    file_size = uploaded_file.size

    if file_size > MAX_FILE_BYTES:
        file_exceeded = True
        st.error(f"Uploaded file exceeds {MAX_FILE_MB} MB limit.")

    else:
        if uploaded_file.name.endswith(".docx"):
            file_text = "\n".join([p.text for p in docx.Document(uploaded_file).paragraphs])

        else:
            file_text = str(uploaded_file.read(), "utf-8")
            uploaded_file.seek(0)

        st.session_state.process_context = file_text
        process_context = file_text
# -------- USE MANUAL INPUT --------
elif manual_input and not manual_exceeded:
            st.session_state.process_context = manual_input
            process_context = manual_input
generate_disabled = (
    (not process_context) or
    manual_exceeded or
    file_exceeded
)

if st.button("Generate Process Design Document", disabled=generate_disabled):
    branding, contacts, company, clients = load_config()
    if not process_context:
        st.error("Please provide process details.")
    else:
        with st.spinner("Generating document..."):
            short_context = get_short_context(client, process_context)   
            today = date.today().strftime("%m/%d/%Y")
            if uploaded_file:
                raw_filename = uploaded_file.name.rsplit('.', 1)[0]
                dynamic_title = raw_filename.replace('_', ' ').replace('-', ' ').title()
                client_name = dynamic_title.split(' ')[0]
                client_code = client_name.upper()
                client_cfg = clients.get(client_code)
                if client_cfg is None:
                    client_cfg = list(clients.values())[0]
                brand = branding.get(client_code, {})

                logo_file = brand.get("Logo", "KMG_LOGO.png")
                banner_file = brand.get("Banner", "KMG_BANNER.png")

                BASE_DIR = os.path.dirname(os.path.abspath(__file__))

                logo_path = os.path.join(BASE_DIR, "Assets", logo_file)
                banner_path = os.path.join(BASE_DIR, "Assets", banner_file)
            else:
                #create dynamic title from manual input
                first_line = manual_input.strip().split("\n")[0]

                if first_line:
                    dynamic_title = first_line[:60]
                else:
                    dynamic_title = "Process Design Document"

                client_name = dynamic_title.split(" ")[0].upper()
                client_code = client_name
                client_cfg = clients.get(client_name)

                if client_cfg is None:
                    client_cfg = list(clients.values())[0]     
                brand = branding.get(client_code, {})
                logo_file = brand.get("Logo", "KMG_LOGO.png")
                banner_file = brand.get("Banner", "KMG_BANNER.png")
                logo_path = os.path.join("Assets", logo_file)
                banner_path = os.path.join("Assets", banner_file)  

            doc = Document()
            section = doc.sections[0]
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            for section in doc.sections:
                section.header_distance = Inches(0.6)
            for sec in doc.sections:
                sec.footer_distance = Inches(0.5)
            from docx.enum.text import WD_LINE_SPACING
            style = doc.styles['Normal']
            style.font.name = 'Trebuchet MS'
            style.font.size = Pt(10)
            pformat = style.paragraph_format
            pformat.space_before = Pt(0)
            pformat.space_after = Pt(0)
            pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pformat.line_spacing = 1

                                                                    # ==============================
                                                                    # PAGE 1: COVER PAGE
                                                                    # ==============================
            
            cover_table = doc.add_table(rows=1, cols=1)
            cover_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_table.autofit = False
            cover_table.columns[0].width = Inches(6.2)
            cover_cell = cover_table.cell(0, 0)
            p_title = cover_cell.add_paragraph()
            p_title.paragraph_format.space_before = Pt(20)
            set_font(p_title.add_run(dynamic_title), size=28, color=KMG_NAVY, bold=True)
            p_title.paragraph_format.space_after = Pt(12)
            # Subtitle → Process Flow Document
            p_sub = cover_cell.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(2)
            run_sub = p_sub.add_run("Process Flow Document")
            set_font(
                run_sub,
                name="Trebuchet MS",
                size=14,
                color=BLACK,
                bold=True
            )
            # Description → Document describing the process of xyz
            p_desc = cover_cell.add_paragraph()
            p_desc.paragraph_format.space_before = Pt(0)
            run_desc = p_desc.add_run(f"Document describing the process of {dynamic_title}")
            set_font(
                run_desc,
                name="Trebuchet MS",
                size=10,
                color=TEXT_GREY,
                italic=True
            )
            p_desc.paragraph_format.space_after = Pt(60)

            def add_meta_line_bold(label, value):
                p = cover_cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(f"{label} {value}")
                set_font(run, name="Trebuchet MS", size=9, bold=True, color=BLACK)
            add_meta_line_bold("Ref #", "Process Flow Document")
            add_meta_line_bold("Date:", today)
            kmg = contacts.get("KMG_CONTACT", {})
            add_meta_line_bold(
                "KMG Contact:",
                f"{kmg.get('Name','')} | {kmg.get('Email','')}"
                            )
            cover_cell.add_paragraph().paragraph_format.space_after = Pt(8)
            # Calculate usable width (once)
            section = doc.sections[0]
            page_width = section.page_width
            left_margin = section.left_margin
            right_margin = section.right_margin
            usable_width = page_width - left_margin - right_margin
            # Top blue line
            line_top = cover_cell.add_paragraph()
            line_top.paragraph_format.space_after = Pt(4)
            run = line_top.add_run(" ")
            run.font.size = Pt(1)
            p = line_top._p
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '24')
            bottom.set(qn('w:color'), '1F497D')
            pBdr.append(bottom)
            pPr.append(pBdr)

            if os.path.exists(banner_path):
                p_img = cover_cell.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.add_run().add_picture(banner_path, width=Inches(7))
            # Bottom blue line
            line_bottom = cover_cell.add_paragraph()
            line_bottom.paragraph_format.space_before = Pt(4)
            run = line_bottom.add_run(" ")
            run.font.size = Pt(1)
            p = line_bottom._p
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '24')
            bottom.set(qn('w:color'), '1F497D')
            pBdr.append(bottom)
            pPr.append(pBdr)
            cover_cell.add_paragraph()
            footer_tab = cover_cell.add_table(1, 2)
            footer_tab.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_tab.autofit = False
            from docx.shared import Emu
            total_width = usable_width
            col1_width = Emu(int(total_width * 0.7))
            col2_width = Emu(int(total_width * 0.3))
            footer_tab.columns[0].width = col1_width
            footer_tab.columns[1].width = col2_width
            footer_tab.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            footer_tab.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell = footer_tab.cell(0, 0)
            # Company name
            p1 = cell.paragraphs[0]
            run_company = p1.add_run(company["Company Name"])
            set_font(run_company, name="Trebuchet MS", size=12, bold=True)
            p1.paragraph_format.space_after = Pt(6)
            # Address
            p2 = cell.add_paragraph()
            run_addr = p2.add_run(company["address"])
            set_font(run_addr, name="Trebuchet MS", size=10)
            # Phone + Fax
            p3 = cell.add_paragraph()
            run_contact = p3.add_run(f"Ph: {company['phone']} | Fax: {company['fax']}")
            set_font(run_contact, name="Trebuchet MS", size=10)
            # Website + social
            p4 = cell.add_paragraph()
            run_web = p4.add_run(f"{company['website']} | {company['social']}")
            set_font(run_web, name="Trebuchet MS", size=10)
            if os.path.exists(logo_path):

                c_right = footer_tab.cell(0, 1).paragraphs[0]
                c_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                c_right.add_run().add_picture(logo_path, width=Inches(1.1))

                                                                    # ==============================
                                                                    # PAGE 2: TABLE OF CONTENTS
                                                                    # ==============================

            cover_cell.add_paragraph()
            doc.sections[0].different_first_page_header_footer = False
            from docx.enum.section import WD_SECTION
            doc.add_section(WD_SECTION.NEW_PAGE)
            doc.sections[-1].header.is_linked_to_previous = False
            insert_constant_header(doc, dynamic_title, client_name, today, logo_path, client_cfg)
            p = doc.add_paragraph()
            run = p.add_run("CONTENTS")
            set_font(run, name="Trebuchet MS", size=18, bold=True, underline=True, color=BLACK)
            p.paragraph_format.space_after = Pt(12)
            contents = [("1 Version History", "3"), ("   1.1 Release History", "3"), ("   1.2 Contact Information", "3"),

                        ("2 Introduction", "4"), ("3 Audience", "4"), ("4 Purpose", "5"), ("5 Scope", "5"), ("6 Process Flow Diagram", "6")]
            for item, pg in contents:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)  
                p.paragraph_format.space_after = Pt(0)
                if item.strip().startswith("1.1") or item.strip().startswith("1.2"):
                    p.paragraph_format.left_indent = Inches(0.25)
                dots = "." * (95 - len(item))
                run = p.add_run(f"{item} {dots} {pg}")
                set_font(run, name="Trebuchet MS", size=11, color=BLACK)
            # --- Page Numbering ---
            section = doc.sections[-1]
            footer = section.footer
            footer.is_linked_to_previous = False
            foot_p = footer.paragraphs[0]
            foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            def create_field(parent_run, field_text):

                fldChar_begin = OxmlElement('w:fldChar'); fldChar_begin.set(qn('w:fldCharType'), 'begin')
                parent_run._r.append(fldChar_begin)
                instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = field_text
                parent_run._r.append(instrText)
                fldChar_end = OxmlElement('w:fldChar'); fldChar_end.set(qn('w:fldCharType'), 'end')
                parent_run._r.append(fldChar_end)
            run_pg = foot_p.add_run()
            create_field(run_pg, "PAGE")
            run_pg.add_text(" of ")
            create_field(run_pg, "NUMPAGES")
            set_font(foot_p.add_run(" | All Rights to this Document Reserved with Key Management Group, Inc."), size=11, color=TEXT_GREY)

                                                                    # ==============================
                                                                    # PAGE 3: VERSION HISTORY
                                                                    # ==============================

            doc.add_page_break()
            vh_p = doc.add_paragraph()
            vh_p.paragraph_format.space_after = Pt(12)
            set_font(vh_p.add_run("VERSION HISTORY"), size=18, bold=True, underline=True)
            vh_p.paragraph_format.space_after = Pt(14)
            p_rh = doc.add_paragraph()
            set_font(p_rh.add_run("1.1  RELEASE HISTORY"), size=14, bold=True)
            p_rh.paragraph_format.space_after = Pt(6)
            rel_tab = doc.add_table(rows=5, cols=6); rel_tab.style = 'Table Grid'
            for i, h in enumerate(["Version", "Date", "Description", "Reason", "Author(s)", "Reviewer"]):
                set_font(rel_tab.cell(0, i).paragraphs[0].add_run(h), size=9, bold=True)
            r1 = rel_tab.rows[1].cells
            row_cells = rel_tab.rows[1].cells
            author = contacts.get("AUTHOR", {})
            data = [
                "0.1",
                today,
                "Initial Draft",
                f"Process Design for {dynamic_title}",
                
                f"{author.get('Name','').replace(' ', chr(10))}",
                ""
            ]
            for i, val in enumerate(data):
                p = row_cells[i].paragraphs[0]
                run = p.add_run(val)
                set_font(run, size=9)  
            gap = doc.add_paragraph()
            gap.paragraph_format.space_after = Pt(10)
            p_contact = doc.add_paragraph()
            p_contact.paragraph_format.space_before = Pt(4)
            p_contact.paragraph_format.space_after = Pt(6)
            set_font(p_contact.add_run("1.2  CONTACT INFORMATION"), size=12, bold=True)
            p_contact_text = doc.add_paragraph()
            run = p_contact_text.add_run(
                f"{contacts['KMG_CONTACT']['Name']} | "
                f"{contacts['KMG_CONTACT']['Title']} | "
                f"{contacts['KMG_CONTACT']['Email']}\n"
            )
            set_font(run, size=9)
            p_contact_text.paragraph_format.space_after = Pt(10)
            p_comp = doc.add_paragraph()
            p_comp.paragraph_format.space_before = Pt(6)
            run = p_comp.add_run("Company Information,")
            set_font(run, size=10, bold=True, italic=True)
            run = p_comp.add_run("\nKey Management Group, Inc.")
            set_font(run, size=9)
            run = p_comp.add_run("\n420 Jericho Turnpike, Suite #215, Jericho. NY - 11753")
            set_font(run, size=9)
            run = p_comp.add_run("\nwww.kmgus.com | 631-777-2424 (phone) | 631-777-2626 (fax)")
            set_font(run, size=9)

                                                                    # ==============================
                                                                    # PAGE 4: AI CONTENT
                                                                    # ==============================

            doc.add_page_break()
            top_gap = doc.add_paragraph()
            top_gap.paragraph_format.space_after = Pt(12)
            sections = ["INTRODUCTION", "AUDIENCE", "PURPOSE", "SCOPE"]
            for title in sections:
                # HEADING
                p_head = doc.add_paragraph()
                p_head.paragraph_format.space_before = Pt(12)
                p_head.paragraph_format.space_after = Pt(6)
                set_font(p_head.add_run(title), size=18, bold=True, underline=True)
                # CONTENT PARAGRAPH
                content = generate_ai_content(client, title, short_context, dynamic_title)

                for line in content.split("\n"):
                    line = line.strip()

                    clean = line.lstrip("*-• ").strip()

                    # detect list item by context (multi-line list from AI)
                    if (
                        clean
                        and len(clean.split()) < 12
                        and line[0].islower() is False
                        and not clean.endswith(".")
                    ):
                        bullet = doc.add_paragraph(clean, style="List Bullet")
                        set_font(bullet.runs[0], size=11)

                    elif clean:
                        para = doc.add_paragraph(clean)
                        for run in para.runs:
                            set_font(run, size=11)
                    elif line:
                        para = doc.add_paragraph(line)
                        for run in para.runs:
                             set_font(run, size=11)

                                                                    # ==============================
                                                                    # PAGE 5: PROCESS FLOW
                                                                    # ==============================

            doc.add_page_break()
            p_flow = doc.add_paragraph()
            p_flow.paragraph_format.space_after = Pt(12)
            set_font(p_flow.add_run("PROCESS FLOW DIAGRAM"), size=18, bold=True, underline=True)
            p_flow.paragraph_format.space_after = Pt(12)

            try:
                # 1. AI se structured JSON data mangwayein
                flow_data = get_smart_flow_data(client, short_context)
                # 2. Smart Flowchart image generate karein (Naya function name)
                chart_filename = generate_smart_flowchart(flow_data)
                # 3. Image ko Document mein insert karein
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(chart_filename, width=Inches(5.0)) 
            except Exception as e:
                doc.add_paragraph(f"Could not generate flowchart: {str(e)}")
            fname = f"KMG_PDD_{dynamic_title.replace(' ', '_')}.docx"
            doc.save(fname)
            st.success("Process Design Document Generated!")
            with open(fname, "rb") as f:
                st.download_button("Download Process Design Document", f, file_name=fname)
